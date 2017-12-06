#!/usr/bin/env python
from docx import Document

document = Document()

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')

prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')

document.add_heading('The REAL meaning of the universe')
document.add_heading('The role of dolphins', level=2)
document.add_page_break()

# Save the document
document.save('hellowrld.docx')

